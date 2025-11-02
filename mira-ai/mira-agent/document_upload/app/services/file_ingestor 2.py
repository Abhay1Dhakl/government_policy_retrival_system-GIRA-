import mimetypes
import re
import uuid
from datetime import datetime
from typing import List, Dict, Tuple
from document_upload.app.models.document import store_document, get_embedding_data
from PyPDF2 import PdfReader
from docx import Document
import io
import hashlib

try:
    from sentence_transformers import SentenceTransformer
    from nltk.tokenize import sent_tokenize
    import nltk
    # NLTK data should be pre-installed in Docker image
    # Check if punkt tokenizer is available
    try:
        nltk.data.find('tokenizers/punkt')
        print("NLTK punkt tokenizer available")
    except LookupError:
        print("Warning: NLTK punkt tokenizer not found. Sentence splitting may not work correctly.")
        # Provide fallback for sent_tokenize
        sent_tokenize = lambda text: [text]  # Fallback: treat whole text as one sentence
    SEMANTIC_CHUNKING_AVAILABLE = True
except ImportError as e:
    SEMANTIC_CHUNKING_AVAILABLE = False
    print(f"Warning: sentence_transformers or nltk not available. Using basic chunking. Error: {e}")
    # Provide fallback for sent_tokenize
    sent_tokenize = lambda text: [text]

_SENTENCE_MODEL = None

# Optional PDF text/table extraction
try:
    import pdfplumber
    PDF_TABLES_AVAILABLE = True
    print("pdfplumber available: using high-fidelity PDF text extraction when possible")
except ImportError as _pdfplumber_err:
    PDF_TABLES_AVAILABLE = False
    print(f"Warning: pdfplumber not installed. Table/text extraction disabled. Reason: {_pdfplumber_err}")

# Optional Camelot table extraction
try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    print("Warning: camelot not installed. Advanced table extraction disabled.")


def get_sentence_model():
    """Return a cached sentence transformer instance."""
    global _SENTENCE_MODEL
    if not SEMANTIC_CHUNKING_AVAILABLE:
        return None
    if _SENTENCE_MODEL is None:
        try:
            _SENTENCE_MODEL = SentenceTransformer("BAAI/bge-m3")
        except Exception as exc:
            print(f"Warning: Could not initialize sentence transformer: {exc}")
            _SENTENCE_MODEL = None
    return _SENTENCE_MODEL


class ChunkingConfig:
    def __init__(self, chunk_size=950, chunk_overlap=200, min_chunk_size=550, max_chunk_size=1400,
                 method="semantic", overlap_sentences=2):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.method = method
        self.overlap_sentences = overlap_sentences

class HealthcareChunker:
    def __init__(self):
        self.config = ChunkingConfig(
            chunk_size=950,
            chunk_overlap=200,
            min_chunk_size=550,
            max_chunk_size=1400,
            method="semantic",
            overlap_sentences=3
        )

        # Initialize sentence transformer once for the process
        self.sentence_model = get_sentence_model()
    
    def healthcare_semantic_chunking(self, text: str, page_number: int = 1) -> List[Dict]:
        """
        Healthcare-optimized chunking that preserves medical context and enables exact citations
        """
        if not SEMANTIC_CHUNKING_AVAILABLE:
            return self._fallback_chunking(text, page_number)
        
        normalized_text = self._normalize_text(text)
        chunks = []
        
        # Split by sections first (medical documents have clear structure)
        sections = self._extract_medical_sections(normalized_text)
        
        for section_idx, section in enumerate(sections):
            # Process each section separately to maintain context
            section_chunks = self._chunk_medical_section(
                section['content'], 
                section['title'],
                section_idx,
                page_number
            )
            chunks.extend(section_chunks)
        
        return self._merge_low_content_chunks(chunks)

    def _normalize_text(self, text: str) -> str:
        """Normalize whitespace and preserve paragraph structure."""
        if not text:
            return ""
        normalized = text.replace('\r\n', '\n')
        normalized = re.sub(r'[ \t]+', ' ', normalized)
        normalized = re.sub(r'\n{3,}', '\n\n', normalized)
        return normalized.strip()
    
    def _extract_medical_sections(self, text: str) -> List[Dict]:
        """Extract medical document sections with robust header detection."""
        sections: List[Dict] = []
        current_section = {"title": "DOCUMENT_START", "content": "", "start_line": 0}
        lines = text.split('\n')

        for idx, line in enumerate(lines):
            stripped = line.strip()
            if self._looks_like_section_header(stripped, lines, idx):
                if current_section["content"].strip():
                    sections.append(current_section.copy())
                current_section = {
                    "title": stripped,
                    "content": "",
                    "start_line": idx + 1
                }
                continue

            current_section["content"] += (stripped + "\n") if stripped else "\n"

        if current_section["content"].strip():
            sections.append(current_section)

        return sections

    def _looks_like_section_header(self, line: str, lines: List[str], index: int) -> bool:
        """Heuristic to determine whether a line is a section header."""
        if not line or len(line) > 120:
            return False

        # Accept numbered headings like "8", "8.4", "2 DOSAGE AND ADMINISTRATION", "8.4 Pediatric Use"
        numbered_heading = re.match(r"^\s*\d+(?:\.\d+)*\s*(?:[\-–:])?\s*[A-Za-z]", line)
        if numbered_heading:
            return True

        if len(re.sub(r'[^A-Za-z]', '', line)) < 4:
            return False
        uppercase_ratio = sum(1 for ch in line if ch.isupper()) / max(len(line), 1)
        if uppercase_ratio < 0.6:
            return False
        # Require following line break or short underline
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if next_line and uppercase_ratio < 0.9 and len(next_line) > 0:
            if len(next_line) > 3 and next_line[0].isalpha():
                return False
        return True
    
    def _chunk_medical_section(self, text: str, section_title: str,
                               section_idx: int, page_number: int) -> List[Dict]:
        """Chunk individual medical sections with line-level tracking"""
        chunks: List[Dict] = []

        try:
            if SEMANTIC_CHUNKING_AVAILABLE:
                sentences = sent_tokenize(text)
            else:
                sentences = re.split(r'(?<=[.!?])\s+', text)
        except Exception:
            sentences = re.split(r'(?<=[.!?])\s+', text)

        sentences = [s.strip() for s in sentences if s.strip()]
        if not sentences:
            return []

        current_chunk: List[str] = []
        current_char_len = 0
        sentence_cursor = 0
        chunk_start_idx = 0

        for sentence in sentences:
            sentence_length = len(sentence)

            if current_chunk and (
                current_char_len + sentence_length > self.config.chunk_size
                and current_char_len >= self.config.min_chunk_size
            ):
                chunk_text = " ".join(current_chunk)
                chunks.append(self._create_medical_chunk(
                    chunk_text,
                    section_title,
                    section_idx,
                    len(chunks),
                    page_number,
                    chunk_start_idx,
                    sentence_cursor - 1
                ))

                overlap_count = min(len(current_chunk), self.config.overlap_sentences)
                current_chunk = current_chunk[-overlap_count:] if overlap_count else []
                current_char_len = sum(len(s) for s in current_chunk)
                chunk_start_idx = max(sentence_cursor - overlap_count, 0)

            if not current_chunk:
                chunk_start_idx = sentence_cursor

            current_chunk.append(sentence)
            current_char_len += sentence_length
            sentence_cursor += 1

            if current_char_len >= self.config.max_chunk_size:
                chunk_text = " ".join(current_chunk)
                chunks.append(self._create_medical_chunk(
                    chunk_text,
                    section_title,
                    section_idx,
                    len(chunks),
                    page_number,
                    chunk_start_idx,
                    sentence_cursor - 1
                ))

                overlap_count = min(len(current_chunk), self.config.overlap_sentences)
                current_chunk = current_chunk[-overlap_count:] if overlap_count else []
                current_char_len = sum(len(s) for s in current_chunk)
                chunk_start_idx = max(sentence_cursor - overlap_count, 0)

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(self._create_medical_chunk(
                chunk_text,
                section_title,
                section_idx,
                len(chunks),
                page_number,
                chunk_start_idx,
                sentence_cursor - 1
            ))

        return chunks
    
    def _create_medical_chunk(self, text: str, section_title: str, section_idx: int,
                              chunk_idx: int, page_number: int,
                              start_sentence_index: int, end_sentence_index: int) -> Dict:
        """Create a medical chunk with enhanced citation metadata"""

        clean_text = text.strip()
        char_count = len(clean_text)
        word_count = len(clean_text.split())
        chunk_type = self._classify_medical_content(clean_text)
        medical_terms_present = self._detect_medical_terms(clean_text)
        relevance_score = self._score_chunk(clean_text, chunk_type)

        return {
            "text": clean_text,
            "section_title": section_title,
            "section_index": section_idx,
            "chunk_index": chunk_idx,
            "page": page_number,
            "word_count": word_count,
            "char_count": char_count,
            "contains_medical_terms": medical_terms_present,
            "chunk_type": chunk_type,
            "has_overlap": chunk_idx > 0,
            "start_sentence_index": start_sentence_index,
            "end_sentence_index": end_sentence_index,
            "relevance_score": relevance_score
        }

    def _score_chunk(self, text: str, chunk_type: str) -> int:
        """Assign a heuristic relevance score to a chunk."""
        keywords = {
            "qt": 6,
            "torsades": 6,
            "cardiovascular": 5,
            "mortality": 5,
            "elderly": 4,
            "contraindication": 4,
            "warning": 3,
            "risk": 3,
            "dosage": 2,
            "interaction": 2
        }

        score = 0
        lowered = text.lower()
        for term, weight in keywords.items():
            if term in lowered:
                score += weight

        if chunk_type in {"safety", "cardiac_safety", "contraindications"}:
            score += 4

        # Reward quantitative data
        if re.search(r"\d+\s*(?:%|per\s*million|mg)", lowered):
            score += 3

        return score

    def _merge_low_content_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Merge adjacent chunks that are too small to stand alone."""
        if not chunks:
            return []

        merged: List[Dict] = []
        for chunk in chunks:
            if merged and chunk["char_count"] < self.config.min_chunk_size:
                previous = merged[-1]
                previous_text = f"{previous['text']} {chunk['text']}".strip()
                previous['text'] = previous_text
                previous['char_count'] = len(previous_text)
                previous['word_count'] = len(previous_text.split())
                previous['end_sentence_index'] = max(previous['end_sentence_index'], chunk['end_sentence_index'])
                previous['relevance_score'] = max(previous['relevance_score'], chunk['relevance_score'])
            else:
                merged.append(chunk)

        for idx, chunk in enumerate(merged):
            chunk['chunk_index'] = idx

        return merged
    
    def _detect_medical_terms(self, text: str) -> bool:
        """Detect if chunk contains medical terminology"""
        medical_indicators = [
            'mg', 'ml', 'diagnosis', 'treatment', 'symptoms', 'patient', 
            'medication', 'dose', 'allergy', 'test', 'result', 'procedure',
            'blood', 'pressure', 'heart', 'lung', 'liver', 'kidney',
            'injection', 'tablet', 'capsule', 'prescription', 'adverse',
            'contraindication', 'indication', 'dosage', 'therapeutic'
        ]
        text_lower = text.lower()
        return any(term in text_lower for term in medical_indicators)

    def _is_pediatric_content(self, text: str, section_title: str = "") -> bool:
        """Detect if chunk relates to pediatric use/populations."""
        hay = f"{section_title}\n{text}".lower()
        pediatric_terms = [
            'pediatric', 'paediatric', 'children', 'child', 'infant', 'neonate', 'neonatal', 'adolescent',
            'under 16', 'under sixteen', 'less than 16', 'age', 'years of age', 'months of age'
        ]
        # Also detect common section codes like 8.4 Pediatric Use
        if re.search(r"^\s*8(\.\d+)*\s+", section_title or '', flags=re.IGNORECASE):
            return True
        return any(term in hay for term in pediatric_terms)
    
    def _classify_medical_content(self, text: str) -> str:
        """Classify medical content type"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['qt interval', 'torsades', 'cardiac', 'cardiovascular']):
            return "cardiac_safety"
        if any(word in text_lower for word in ['diagnosis', 'assessment', 'impression']):
            return "diagnosis"
        elif any(word in text_lower for word in ['medication', 'prescription', 'dose', 'mg']):
            return "medication"
        elif any(word in text_lower for word in ['allergy', 'reaction', 'adverse']):
            return "allergy"
        elif any(word in text_lower for word in ['lab', 'test', 'result', 'blood']):
            return "lab_result"
        elif any(word in text_lower for word in ['procedure', 'surgery', 'operation']):
            return "procedure"
        elif any(word in text_lower for word in ['contraindication', 'warning', 'precaution']):
            return "safety"
        else:
            return "general"
    
    def _fallback_chunking(self, text: str, page_number: int) -> List[Dict]:
        """Fallback chunking when semantic libraries are not available"""
        chunks = []
        
        # Simple paragraph-based chunking
        paragraphs = text.split('\n\n')
        current_chunk = ""
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            if len(current_chunk) + len(paragraph) > self.config.chunk_size and current_chunk:
                chunks.append({
                    "text": current_chunk.strip(),
                    "section_title": "DOCUMENT_CONTENT",
                    "section_index": 0,
                    "chunk_index": chunk_index,
                    "page": page_number,
                    "word_count": len(current_chunk.split()),
                    "char_count": len(current_chunk),
                    "contains_medical_terms": self._detect_medical_terms(current_chunk),
                    "chunk_type": self._classify_medical_content(current_chunk),
                    "has_overlap": chunk_index > 0
                })
                chunk_index += 1
                current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                "text": current_chunk.strip(),
                "section_title": "DOCUMENT_CONTENT",
                "section_index": 0,
                "chunk_index": chunk_index,
                "page": page_number,
                "word_count": len(current_chunk.split()),
                "char_count": len(current_chunk),
                "contains_medical_terms": self._detect_medical_terms(current_chunk),
                "chunk_type": self._classify_medical_content(current_chunk),
                "has_overlap": chunk_index > 0
            })
        
        return chunks

# Lazy global healthcare chunker
_healthcare_chunker = None

def get_healthcare_chunker():
    global _healthcare_chunker
    if _healthcare_chunker is None:
        # Instantiate the chunker lazily, which will only load the sentence model on demand
        _healthcare_chunker = HealthcareChunker()
    return _healthcare_chunker

class HealthcareDocumentStorage:
    """Enhanced document storage for healthcare-specific chunking with precise citations"""
    
    def __init__(self):
        self.chunker = get_healthcare_chunker()

    @staticmethod
    def _normalize_for_metadata(text: str) -> str:
        if not text:
            return ""
        cleaned = re.sub(r"-\s*\n", "", text)
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip()

    @staticmethod
    def _extract_section_code(section_title: str) -> str:
        if not section_title:
            return ""
        match = re.match(r"^(\d+(?:\.\d+)*)", section_title.strip())
        return match.group(1) if match else ""

    @staticmethod
    def _chunk_keywords(text: str) -> List[str]:
        if not text:
            return []
        tokens = re.findall(r"[A-Za-z][A-Za-z-]{3,}", text.lower())
        stop_words = {
            "with", "from", "this", "that", "have", "been", "their", "there", "which",
            "should", "would", "could", "will", "into", "about", "after", "before",
            "during", "without", "between", "patients", "includes", "including", "because",
            "where", "when", "those", "these", "such", "other", "also", "most", "more"
        }
        filtered = [token for token in tokens if token not in stop_words]
        counts = {}
        for token in filtered:
            counts[token] = counts.get(token, 0) + 1
        sorted_tokens = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
        return [token for token, _ in sorted_tokens[:8]]

    @staticmethod
    def _contains_numeric_data(text: str) -> bool:
        if not text:
            return False
        return bool(re.search(r"\b\d+(?:\.\d+)?\s?(?:mg|ml|mcg|g|kg|%)?\b", text.lower()))
    
    def process_document_text(self, text: str, document_name: str, page_number: int = 1) -> List[Dict]:
        """Process document using healthcare-specific chunking"""
        try:
            chunks = self.chunker.healthcare_semantic_chunking(text, page_number)
            
            # Enhance chunks with document metadata
            for chunk in chunks:
                normalized_text = self._normalize_for_metadata(chunk.get("text", ""))
                section_code = self._extract_section_code(chunk.get("section_title", ""))
                sentence_count = 0
                if isinstance(chunk.get("start_sentence_index"), int) and isinstance(chunk.get("end_sentence_index"), int):
                    sentence_count = max(0, chunk["end_sentence_index"] - chunk["start_sentence_index"] + 1)
                keyword_terms = self._chunk_keywords(normalized_text)
                has_numeric = self._contains_numeric_data(normalized_text)

                chunk.update({
                    "document_name": document_name,
                    "citation_id": f"{page_number}.{chunk['chunk_index'] + 1}",
                    "full_citation": f"[{page_number}.{chunk['chunk_index'] + 1}]",
                    "normalized_text": normalized_text,
                    "section_code": section_code,
                    "sentence_count": sentence_count,
                    "keyword_terms": keyword_terms,
                    "has_numeric_data": has_numeric,
                    "is_pediatric": self._is_pediatric_content(cleaned := chunk.get("text", ""), chunk.get("section_title", "")),
                    "metadata": {
                        "source": document_name,
                        "page": page_number,
                        "section": chunk.get("section_title", "UNKNOWN"),
                        "section_code": section_code,
                        "medical_content": chunk.get("chunk_type", "general"),
                        "has_medical_terms": chunk.get("contains_medical_terms", False),
                        "has_numeric_data": has_numeric,
                        "is_pediatric": self._is_pediatric_content(chunk.get("text", ""), chunk.get("section_title", "")),
                        "processing_method": "healthcare_semantic",
                        "relevance_score": chunk.get("relevance_score", 0),
                        "start_sentence_index": chunk.get("start_sentence_index", 0),
                        "end_sentence_index": chunk.get("end_sentence_index", 0),
                        "sentence_count": sentence_count,
                        "keyword_terms": keyword_terms
                    }
                })
            
            return chunks
            
        except Exception as e:
            print(f"Error in healthcare chunking, falling back to standard: {e}")
            return self._fallback_processing(text, document_name, page_number)
    
    def _fallback_processing(self, text: str, document_name: str, page_number: int) -> List[Dict]:
        """Fallback to standard chunking if healthcare chunking fails"""
        standard_chunks = split_text_into_chunks_with_overlap(text, 1200, 250)
        
        processed_chunks = []
        for idx, chunk_text in enumerate(standard_chunks):
            clean_text = chunk_text.strip()
            chunk_type = self.chunker._classify_medical_content(clean_text)
            has_terms = self.chunker._detect_medical_terms(clean_text)
            relevance = self.chunker._score_chunk(clean_text, chunk_type)
            normalized_text = self._normalize_for_metadata(clean_text)
            keyword_terms = self._chunk_keywords(normalized_text)
            has_numeric = self._contains_numeric_data(normalized_text)

            processed_chunks.append({
                "text": clean_text,
                "document_name": document_name,
                "page": page_number,
                "chunk_index": idx,
                "citation_id": f"{page_number}.{idx + 1}",
                "full_citation": f"[{page_number}.{idx + 1}]",
                "section_title": "DOCUMENT_CONTENT",
                "section_index": 0,
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text),
                "contains_medical_terms": has_terms,
                "chunk_type": chunk_type,
                "has_overlap": idx > 0,
                "normalized_text": normalized_text,
                "section_code": "",
                "sentence_count": 0,
                "keyword_terms": keyword_terms,
                "has_numeric_data": has_numeric,
                "is_pediatric": self.chunker._is_pediatric_content(clean_text, "DOCUMENT_CONTENT"),
                "start_sentence_index": 0,
                "end_sentence_index": 0,
                "relevance_score": relevance,
                "metadata": {
                    "source": document_name,
                    "page": page_number,
                    "section": "DOCUMENT_CONTENT",
                    "section_code": "",
                    "medical_content": chunk_type,
                    "has_medical_terms": has_terms,
                    "has_numeric_data": has_numeric,
                    "is_pediatric": self.chunker._is_pediatric_content(clean_text, "DOCUMENT_CONTENT"),
                    "processing_method": "standard_fallback",
                    "relevance_score": relevance,
                    "start_sentence_index": 0,
                    "end_sentence_index": 0,
                    "sentence_count": 0,
                    "keyword_terms": keyword_terms
                }
            })
        
        return processed_chunks

def _hash_table_rows(rows: List[List[str]]) -> str:
    try:
        m = hashlib.sha256()
        preview = "\n".join([",".join([(c or "").strip() for c in r]) for r in rows[:5]])
        m.update(preview.encode("utf-8", errors="ignore"))
        return m.hexdigest()[:16]
    except Exception:
        return str(uuid.uuid4())[:16]

def _table_rows_to_text(rows: List[List[str]], max_rows: int = 30, max_chars: int = 4000) -> str:
    """Convert table rows to a compact CSV-like text for embedding with limits."""
    safe_rows = []
    for r in rows[:max_rows]:
        safe_cells = [(c if isinstance(c, str) else ("" if c is None else str(c))) for c in r]
        safe_rows.append(", ".join(cell.strip() for cell in safe_cells))
    text = "\n".join(safe_rows)
    return text[:max_chars]

def extract_tables_from_pdf_bytes(pdf_bytes: bytes) -> Dict[int, List[Dict]]:
    """Extract tables per page using pdfplumber. Returns {page_number: [table_dict,...]}"""
    results: Dict[int, List[Dict]] = {}
    if not PDF_TABLES_AVAILABLE:
        return results
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                page_tables = []
                for t_idx, rows in enumerate(tables):
                    if not rows or all(not any(cell for cell in row) for row in rows):
                        continue
                    text_repr = _table_rows_to_text(rows)
                    table_chunk = {
                        "text": text_repr,
                        "section_title": "TABLE",
                        "section_index": 0,
                        "chunk_index": t_idx,  # will be normalized later
                        "page": i,
                        "word_count": len(text_repr.split()),
                        "char_count": len(text_repr),
                        "contains_medical_terms": True,  # tables often contain numeric/medical terms
                        "chunk_type": "table",
                        "has_overlap": False,
                        "normalized_text": text_repr,
                        "section_code": "",
                        "sentence_count": 0,
                        "keyword_terms": [],
                        "relevance_score": 0,
                        "is_table": True,
                        "table_rows": len(rows),
                        "table_cols": max((len(r) for r in rows if r), default=0),
                        "table_hash": _hash_table_rows(rows),
                        "metadata": {
                            "source": "PDF_TABLE",
                            "page": i,
                            "section": "TABLE",
                            "medical_content": "table",
                            "has_medical_terms": True,
                            "has_numeric_data": any(any(ch and any(d.isdigit() for d in str(ch)) for ch in r) for r in rows),
                            "processing_method": "pdfplumber_table",
                            "relevance_score": 0,
                            "sentence_count": 0,
                            "table_rows": len(rows),
                            "table_cols": max((len(r) for r in rows if r), default=0),
                            "table_hash": _hash_table_rows(rows),
                        }
                    }
                    page_tables.append(table_chunk)
                if page_tables:
                    results[i] = page_tables
    except Exception as e:
        print(f"Table extraction failed: {e}")
    return results

def extract_camelot_tables_from_bytes(pdf_bytes: bytes) -> Dict[int, List[Dict]]:
    """Extract tables using Camelot (both lattice and stream) and return {page_no: [table_chunks]}"""
    results: Dict[int, List[Dict]] = {}
    if not CAMELOT_AVAILABLE:
        return results
    import tempfile
    import pandas as pd  # camelot uses pandas tables
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name

        combined = []
        for flavor in ("lattice", "stream"):
            try:
                tables = camelot.read_pdf(tmp_path, pages="all", flavor=flavor)
            except Exception:
                tables = []
            for t in tables:
                try:
                    df = t.df if hasattr(t, "df") else None
                    if df is None or df.empty:
                        continue
                    # Normalize dataframe to rows of strings
                    df = df.fillna("")
                    rows = df.astype(str).values.tolist()
                    if not rows:
                        continue
                    page_no = int(getattr(t, "page", 1))
                    text_repr = _table_rows_to_text(rows)
                    chunk = {
                        "text": text_repr,
                        "section_title": "TABLE",
                        "section_index": 0,
                        "chunk_index": 0,  # will reindex later per page
                        "page": page_no,
                        "word_count": len(text_repr.split()),
                        "char_count": len(text_repr),
                        "contains_medical_terms": True,
                        "chunk_type": "table",
                        "has_overlap": False,
                        "normalized_text": text_repr,
                        "section_code": "",
                        "sentence_count": 0,
                        "keyword_terms": [],
                        "relevance_score": 0,
                        "is_table": True,
                        "table_rows": len(rows),
                        "table_cols": max((len(r) for r in rows if r), default=0),
                        "table_hash": _hash_table_rows(rows),
                        "metadata": {
                            "source": "PDF_TABLE",
                            "page": page_no,
                            "section": "TABLE",
                            "medical_content": "table",
                            "has_medical_terms": True,
                            "has_numeric_data": any(any(ch and any(d.isdigit() for d in str(ch)) for ch in r) for r in rows),
                            "processing_method": f"camelot_{flavor}",
                            "relevance_score": 0,
                            "sentence_count": 0,
                            "table_rows": len(rows),
                            "table_cols": max((len(r) for r in rows if r), default=0),
                            "table_hash": _hash_table_rows(rows),
                        }
                    }
                    combined.append(chunk)
                except Exception:
                    continue
        # Group by page
        for c in combined:
            results.setdefault(c["page"], []).append(c)
    except Exception as e:
        print(f"Camelot extraction failed: {e}")
    finally:
        if tmp_path:
            import os
            try:
                os.remove(tmp_path)
            except Exception:
                pass
    return results

def extract_text_per_page_pdfplumber(pdf_bytes: bytes) -> List[dict]:
    """Use pdfplumber to extract text per page (better layout)"""
    if not PDF_TABLES_AVAILABLE:
        return []
    
    results = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_result = {
                    "page": page_num,
                    "text": "",
                    "extraction_method": "standard",
                    "char_count": 0,
                    "word_count": 0,
                    "issues": []
                }
                
                try:
                    # Method 1: Standard extraction
                    text_standard = page.extract_text() or ""
                    
                    # Method 2: Layout-preserving extraction
                    text_layout = page.extract_text(layout=True) or ""
                    
                    # Method 3: Try with different tolerances
                    text_tolerant = page.extract_text(
                        x_tolerance=3, 
                        y_tolerance=3
                    ) or ""
                    
                    # Choose the best result
                    texts = {
                        "standard": text_standard,
                        "layout": text_layout, 
                        "tolerant": text_tolerant
                    }
                    
                    # Select the longest non-empty text
                    best_method = "standard"
                    best_text = text_standard
                    max_length = len(text_standard)
                    
                    for method, text in texts.items():
                        if len(text) > max_length and text.strip():
                            best_method = method
                            best_text = text
                            max_length = len(text)
                    
                    page_result.update({
                        "text": best_text,
                        "extraction_method": best_method,
                        "char_count": len(best_text),
                        "word_count": len(best_text.split()),
                        "all_methods": {k: len(v) for k, v in texts.items()}
                    })
                    
                    # Check for potential issues
                    if not best_text.strip():
                        page_result["issues"].append("empty_text")
                    elif len(best_text) < 50:
                        page_result["issues"].append("very_short_text")
                    if not best_text.endswith(('.', '!', '?', '\n')):
                        page_result["issues"].append("abrupt_ending")
                        
                except Exception as page_error:
                    page_result.update({
                        "text": "",
                        "issues": [f"extraction_error: {str(page_error)}"]
                    })
                
                results.append(page_result)
                
    except Exception as e:
        print(f"pdfplumber text extraction failed: {e}")
    
    return results

def extract_text_per_page_pdfplumber_words(pdf_bytes: bytes, y_tol: float = 3.0) -> List[str]:
    """Fallback: rebuild page text from pdfplumber words with basic line grouping.
    This often recovers text when page.extract_text() returns empty.
    """
    if not PDF_TABLES_AVAILABLE:
        return []
    pages_text: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    words = page.extract_words(x_tolerance=1, y_tolerance=1) or []
                except Exception:
                    words = []
                if not words:
                    pages_text.append("")
                    continue
                # Sort by vertical then horizontal position
                words.sort(key=lambda w: (w.get("top", 0.0), w.get("x0", 0.0)))
                lines: List[List[dict]] = []
                for w in words:
                    if not lines:
                        lines.append([w])
                        continue
                    last_line = lines[-1]
                    last_top = last_line[-1].get("top", 0.0)
                    if abs(w.get("top", 0.0) - last_top) <= y_tol:
                        last_line.append(w)
                    else:
                        lines.append([w])
                # Build text from lines
                line_texts: List[str] = []
                for ln in lines:
                    ln.sort(key=lambda w: w.get("x0", 0.0))
                    line_texts.append(" ".join([w.get("text", "").strip() for w in ln if w.get("text")]))
                pages_text.append("\n".join([lt for lt in line_texts if lt]))
    except Exception as e:
        print(f"pdfplumber words fallback failed: {e}")
    return pages_text

def extract_text_per_page_pymupdf(pdf_bytes: bytes) -> List[str]:
    """Use PyMuPDF (fitz) to extract text per page; often better than PyPDF2."""
    texts: List[str] = []
    try:
        import fitz  # PyMuPDF
    except Exception as e:
        print(f"PyMuPDF not available: {e}")
        return texts
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_index in range(len(doc)):
            try:
                page = doc.load_page(page_index)
                # 'text' uses layout-aware extraction in recent PyMuPDF
                txt = page.get_text("text") or ""
            except Exception:
                txt = ""
            texts.append(txt)
        doc.close()
    except Exception as e:
        print(f"PyMuPDF text extraction failed: {e}")
    return texts

def ocr_pdf_pages_with_tesseract(pdf_bytes: bytes, dpi: int = 200) -> List[str]:
    """OCR pages when a PDF is image-only. Uses PyMuPDF to render pages and pytesseract to extract text."""
    texts: List[str] = []
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except Exception as e:
        print(f"OCR fallback unavailable (missing deps): {e}")
        return texts
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        for page_index in range(len(doc)):
            try:
                page = doc.load_page(page_index)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                txt = pytesseract.image_to_string(img) or ""
            except Exception as e:
                print(f"OCR failed on page {page_index + 1}: {e}")
                txt = ""
            texts.append(txt)
        doc.close()
    except Exception as e:
        print(f"OCR processing failed: {e}")
    return texts

# Initialize global healthcare document storage
healthcare_storage = HealthcareDocumentStorage()

def split_text_into_chunks_with_overlap(text, chunk_size=1000, overlap_size=200):
    """
    Split text into chunks with overlap to preserve context across boundaries.
    
    Args:
        text (str): The text to split
        chunk_size (int): Target size for each chunk in characters
        overlap_size (int): Number of characters to overlap between chunks
        
    Returns:
        list: List of text chunks with overlap
    """
    if overlap_size >= chunk_size:
        raise ValueError("Overlap size must be less than chunk size")
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        # Define end position for this chunk
        end = start + chunk_size
        
        # If this is the last chunk, take everything remaining
        if end >= text_length:
            chunk = text[start:]
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk.strip())
            break
        
        # Try to find a good breaking point (sentence boundary)
        chunk_text = text[start:end]
        
        # Look for sentence endings within the last 20% of the chunk
        search_start = max(0, len(chunk_text) - int(chunk_size * 0.2))
        sentence_end_pattern = r'[.!?]\s+'
        
        # Find the last sentence ending in the search area
        matches = list(re.finditer(sentence_end_pattern, chunk_text[search_start:]))
        
        if matches:
            # Use the last sentence ending found
            last_match = matches[-1]
            actual_end = start + search_start + last_match.end()
            chunk = text[start:actual_end].strip()
        else:
            # Fall back to word boundary
            words = chunk_text.split()
            if len(words) > 1:
                # Remove the last word to avoid cutting words in half
                chunk = ' '.join(words[:-1])
                actual_end = start + len(chunk)
            else:
                # Single word longer than chunk_size
                chunk = chunk_text
                actual_end = end
        
        if chunk.strip():
            chunks.append(chunk.strip())

        # Move start position with overlap (guarantee forward progress)
        next_start = actual_end - overlap_size
        # Ensure we never go backwards or get stuck
        start = max(next_start, start + 1, actual_end if next_start <= start else next_start)
    
    print(f"Total chunks created: {len(chunks)} (chunk_size: {chunk_size}, overlap: {overlap_size})")
    return chunks

def split_text_into_chunks_advanced(text, chunk_size=1000, overlap_size=200, min_chunk_size=100):
    """
    Advanced chunking with medical document optimization and overlap.
    
    Args:
        text (str): The text to split
        chunk_size (int): Target size for each chunk
        overlap_size (int): Overlap between chunks
        min_chunk_size (int): Minimum acceptable chunk size
        
    Returns:
        list: List of optimized text chunks
    """
    if not text or not text.strip():
        print("Empty or None text provided to split_text_into_chunks_advanced")
        return []
    
    print(f" Chunking text: {len(text)} characters")
    
    # First, try to split by paragraphs
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
            
        # If paragraph is too long, split it
        if len(paragraph) > chunk_size:
            # Save current chunk if it exists
            if current_chunk and len(current_chunk) >= min_chunk_size:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            
            # Split the long paragraph with overlap
            para_chunks = split_text_into_chunks_with_overlap(
                paragraph, chunk_size, overlap_size
            )
            chunks.extend(para_chunks)
            continue
        
        # Check if adding this paragraph would exceed chunk size
        if current_chunk and len(current_chunk) + len(paragraph) + 2 > chunk_size:
            # Save current chunk
            if len(current_chunk) >= min_chunk_size:
                chunks.append(current_chunk.strip())
            
            # Start new chunk with overlap from previous
            if overlap_size > 0 and chunks:
                prev_chunk = chunks[-1]
                overlap_text = prev_chunk[-overlap_size:] if len(prev_chunk) > overlap_size else prev_chunk
                # Find a good breaking point for overlap
                words = overlap_text.split()
                if len(words) > 3:
                    overlap_text = ' '.join(words[-3:])  # Last 3 words
                current_chunk = overlap_text + " " + paragraph
            else:
                current_chunk = paragraph
        else:
            # Add to current chunk
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph
    
    # Add the last chunk
    # Always keep the final chunk if it has any content to avoid truncation
    if current_chunk and current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    print(f"Advanced chunking created: {len(chunks)} chunks (target: {chunk_size}, overlap: {overlap_size})")
    return chunks

# Keep the original function for backward compatibility
def split_text_into_chunks(text, chunk_size=1000):
    """Original chunking function - kept for backward compatibility."""
    return split_text_into_chunks_with_overlap(text, chunk_size, overlap_size=200)

async def ingest_from_file(data, file):
    content = await file.read()
    print(f" File processing - filename: {file.filename}, content length: {len(content)}")
    print(f" Data contains file_name: {data.get('file_name', 'Not provided')}")
    
    # Try to get the original filename from data if available
    original_filename = data.get('file_name', file.filename)
    print(f" Using filename for MIME detection: {original_filename}")
    
    # Try to get mime type from multiple sources
    mime_type, _ = mimetypes.guess_type(original_filename)
    print(f" MIME type from filename: {mime_type}")
    
    # If mime type detection failed, try other methods
    if not mime_type:
        # Check if it's a PDF by content signature
        if content.startswith(b'%PDF'):
            mime_type = "application/pdf"
            print(" Detected PDF from content signature")
        # Check for DOCX signature
        elif content.startswith(b'PK') and b'word/' in content[:1000]:
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            print(" Detected DOCX from content signature")
        else:
            print(" Could not determine file type, assuming text")
            mime_type = "text/plain"
    
    print(f" Final MIME type: {mime_type}")
    text = None
    metadata = []

    if mime_type and mime_type.startswith("text"):
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin1', errors='replace')  # Fallback for unusual encodings

    elif mime_type == "application/pdf":
        try:
            print(" Processing PDF file with PyMuPDF...")
            metadata = []
            chunks = []

            # Use PyMuPDF for text extraction; fallback to pdfplumber, PyPDF2, then OCR
            page_texts = extract_text_per_page_pymupdf(content)
            if not page_texts or all((not t or not t.strip()) for t in page_texts):
                print(" PyMuPDF returned no/empty text; trying pdfplumber next")
                page_texts_dicts = extract_text_per_page_pdfplumber(content)
                if page_texts_dicts and isinstance(page_texts_dicts[0], dict):
                    page_texts = [d.get('text', '') for d in page_texts_dicts]
                else:
                    page_texts = page_texts_dicts

            if not page_texts or all((not t or not t.strip()) for t in page_texts):
                print(" pdfplumber returned no/empty text; trying PyPDF2 next")
                pdf_file = io.BytesIO(content)
                reader = PdfReader(pdf_file)
                page_texts = []
                for page in reader.pages:
                    try:
                        page_texts.append(page.extract_text() or "")
                    except Exception:
                        page_texts.append("")

            if not page_texts or all((not t or not t.strip()) for t in page_texts):
                print(" PyPDF2 returned no/empty text; attempting OCR fallback with Tesseract")
                page_texts = ocr_pdf_pages_with_tesseract(content)

            print(f" PDF has {len(page_texts)} pages")

            for idx, page_text in enumerate(page_texts, start=1):
                try:
                    orig_len = len(page_text) if page_text else 0
                    used_method = "pymupdf"

                    print(f" Page {idx}: extracted {orig_len} characters via {used_method}")
                    if not page_text or not page_text.strip():
                        print(f"Page {idx} is empty or has no extractable text")
                        metadata.append({"page": idx, "text": ""})
                        continue

                    metadata.append({"page": idx, "text": page_text})

                    healthcare_chunks = healthcare_storage.process_document_text(
                        page_text,
                        original_filename,
                        page_number=idx
                    )
                    print(f" Page {idx}: created {len(healthcare_chunks)} healthcare chunks")
                    chunks.extend(healthcare_chunks)
                except Exception as page_error:
                    print(f"Error processing page {idx}: {page_error}")
                    metadata.append({"page": idx, "text": ""})
                    continue

            # Table extraction (optional): You may add PyMuPDF-based or Camelot table extraction here if needed
            # (Current implementation omits pdfplumber table extraction)

            print(f" Total chunks created: {len(chunks)}")

            chunk_metadata_with_data = {
                **data,
                "file_name": original_filename,  # Use the original filename with proper extension
                "mime_type": mime_type,
                "source_metadata": metadata,   # full page texts
                "chunks": chunks,              # chunks with page numbers and overlap info
                "chunking_method": "advanced_overlap",
                "chunk_config": {
                    "chunk_size": 1000,
                    "overlap_size": 200,
                    "min_chunk_size": 100,
                    "tables": False
                }
            }

            await store_document(chunk_metadata_with_data, None)

            return {
                "status": "success", 
                "chunks_stored": len(chunks),
                "chunking_method": "advanced_overlap",
                "total_pages": len(page_texts)
            }

        except Exception as e:
            print(f"Error reading PDF: {e}")
            text = "[PDF content could not be extracted]"

    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = Document(io.BytesIO(content))
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            
            # Use healthcare-specific chunking for DOCX files
            chunks = healthcare_storage.process_document_text(
                full_text,
                original_filename,
                page_number=1  # DOCX files are treated as single "page"
            )
            
            metadata = [{"paragraph_index": i, "text": paragraph.text} for i, paragraph in enumerate(doc.paragraphs) if paragraph.text]
            
            chunk_metadata_with_data = {
                **data,
                "file_name": original_filename,
                "mime_type": mime_type,
                "source_metadata": metadata,
                "chunks": chunks,
                "chunking_method": "healthcare_semantic"
            }
            
            await store_document(chunk_metadata_with_data, None)
            
            return {
                "status": "success",
                "chunks_stored": len(chunks),
                "chunking_method": "advanced_overlap"
            }
            
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            text = "[DOCX content could not be extracted]"
    else:
        text = f"[Unsupported file type: {file.filename}]"

    # Handle text files and fallback cases
    if text:
        # Use healthcare-specific chunking for better medical document processing
        chunks = healthcare_storage.process_document_text(
            text, 
            original_filename,
            page_number=1  # Text files are single "page"
        )
        
        chunk_metadata_with_data = {
            **data,
            "file_name": original_filename,
            "mime_type": mime_type,
            "chunks": chunks,
            "chunking_method": "healthcare_semantic"
        }
        
        await store_document(chunk_metadata_with_data, None)

    print(f"File type: {mime_type}, Content Length: {len(content) if content else 0}")
    return {"status": "success", "chunks_stored": len(chunks) if 'chunks' in locals() else 0}